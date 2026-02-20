# 优化后的简历解析函数
def parse_resume_optimized(file_path):
    """解析PDF/Word简历，提取结构化字段 - 优化版本"""
    print('[DEBUG] parse_resume_optimized called with file_path:', file_path)
    
    # 延迟导入模块
    print('[DEBUG] Importing base modules...')
    try:
        import os
        import re
        import PyPDF2
        
        print('[DEBUG] Base modules imported successfully')
    except Exception as e:
        print('[ERROR] Base module import failed:', str(e))
        import traceback
        traceback.print_exc()
        return None
    print('[DEBUG] Current directory:', os.getcwd())
    print('[DEBUG] File exists:', os.path.exists(file_path))
    print('[DEBUG] File size:', os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A')

    text = ""
    avatar_path = None
    # 解析PDF
    if file_path.endswith('.pdf'):
        try:
            print('[DEBUG] Parsing PDF file:', file_path)
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                print('[DEBUG] PDF has', len(reader.pages), 'pages')
                
                # 遍历每一页
                for page_num, page in enumerate(reader.pages):
                    print('[DEBUG] Processing page', page_num + 1)
                    try:
                        page_text = page.extract_text() or ""
                        print('[DEBUG] Extracted', len(page_text), 'chars from page', page_num + 1)
                        text += page_text
                    except Exception as e:
                        print('[ERROR] Failed to extract text from page', page_num + 1, ':', str(e))
                        import traceback
                        traceback.print_exc()
                        continue
                
                print('[DEBUG] PDF parsing completed successfully')
                print('[DEBUG] Total extracted text length:', len(text))
        except Exception as e:
            print('[ERROR] PDF parsing failed:', str(e))
            import traceback
            traceback.print_exc()
            return None
    # 解析Word（docx）
    elif file_path.endswith('.docx'):
        try:
            print('[DEBUG] Parsing DOCX file:', file_path)
            
            # 方法1：使用python-docx库
            print('[DEBUG] Method 1: Using python-docx library...')
            try:
                from docx import Document
                print('[DEBUG] docx module imported successfully')
                
                doc = Document(file_path)
                print('[DEBUG] DOCX document loaded successfully')
                print('[DEBUG] Found', len(doc.paragraphs), 'paragraphs')
                
                # 提取段落文本
                print('[DEBUG] Extracting paragraphs...')
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + '\n'
                        print('[DEBUG] Found paragraph:', paragraph.text[:50], '...' if len(paragraph.text) > 50 else '')
                
                # 提取表格内容
                print('[DEBUG] Extracting tables...')
                for table in doc.tables:
                    print('[DEBUG] Found table with', len(table.rows), 'rows')
                    for row in table.rows:
                        row_text = ''
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text += cell.text + '\t'
                        if row_text.strip():
                            text += row_text + '\n'
                            print('[DEBUG] Found table row:', row_text[:100], '...' if len(row_text) > 100 else '')
                
                # 提取页眉页脚
                print('[DEBUG] Extracting headers and footers...')
                for section in doc.sections:
                    # 提取页眉
                    if section.header and section.header.paragraphs:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                text += paragraph.text + '\n'
                    # 提取页脚
                    if section.footer and section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + '\n'
            except Exception as e:
                print('[DEBUG] python-docx extraction failed:', str(e))
            
            # 检查是否提取到有效文本
            print('[DEBUG] Text extracted via python-docx:', len(text), 'chars')
            if len(text) < 100 or text.strip() == '0':
                # 方法2：直接从ZIP文件中提取XML
                print('[DEBUG] Method 2: Direct XML extraction from ZIP...')
                try:
                    import zipfile
                    import xml.etree.ElementTree as ET
                    
                    with zipfile.ZipFile(file_path, 'r') as zf:
                        print('[DEBUG] Successfully opened DOCX as ZIP file')
                        
                        # 读取document.xml
                        if 'word/document.xml' in zf.namelist():
                            print('[DEBUG] Found word/document.xml')
                            with zf.open('word/document.xml') as f:
                                xml_content = f.read().decode('utf-8', errors='ignore')
                                print('[DEBUG] XML content length:', len(xml_content), 'chars')
                                
                            # 解析XML
                            root = ET.fromstring(xml_content)
                            
                            # Word XML命名空间
                            ns = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            }
                            
                            # 提取所有文本节点
                            xml_text = []
                            for elem in root.findall('.//w:t', ns):
                                if elem.text:
                                    xml_text.append(elem.text)
                            
                            # 合并文本
                            text = ' '.join(xml_text)
                            print('[DEBUG] Text extracted via XML:', len(text), 'chars')
                            if len(text) > 0:
                                print('[DEBUG] First 200 chars of XML text:', repr(text[:200]))
                        else:
                            print('[DEBUG] word/document.xml not found')
                except Exception as e:
                    print('[DEBUG] XML extraction failed:', str(e))
                    import traceback
                    traceback.print_exc()
            
            # 提取个人照片
            print('[DEBUG] Extracting personal photo...')
            avatar_path = None
            try:
                import zipfile
                import os
                import uuid
                from PIL import Image
                import io
                
                with zipfile.ZipFile(file_path, 'r') as zf:
                    # 查找所有图片文件
                    image_files = []
                    for name in zf.namelist():
                        if name.startswith('word/media/') and (name.endswith('.png') or name.endswith('.jpg') or name.endswith('.jpeg')):
                            image_files.append(name)
                    
                    print('[DEBUG] Found', len(image_files), 'image files')
                    
                    if image_files:
                        # 分析图片，选择最可能是个人照片的图片
                        best_photo = None
                        best_score = 0
                        
                        for img_path in image_files:
                            try:
                                # 读取图片信息
                                with zf.open(img_path) as f:
                                    img_data = f.read()
                                
                                # 分析图片
                                img = Image.open(io.BytesIO(img_data))
                                width, height = img.size
                                aspect_ratio = width / height
                                file_size = len(img_data)
                                
                                print(f'[DEBUG] Image {img_path}: {width}x{height}, aspect={aspect_ratio:.2f}, size={file_size/1024:.1f}KB')
                                
                                # 评分规则
                                score = 0
                                
                                # 1. 个人照片通常是正方形或接近正方形
                                if 0.8 <= aspect_ratio <= 1.2:
                                    score += 5
                                
                                # 2. 个人照片大小适中（10KB-500KB）
                                if 10*1024 <= file_size <= 500*1024:
                                    score += 3
                                
                                # 3. 个人照片分辨率适中（200x200以上，1000x1000以下）
                                if 200 <= width <= 1000 and 200 <= height <= 1000:
                                    score += 3
                                
                                # 4. 排除小标签图片（通常很小）
                                if width < 100 or height < 100:
                                    score -= 5
                                
                                # 5. 排除过大的图片（可能是背景或其他元素）
                                if width > 1000 or height > 1000:
                                    score -= 3
                                
                                # 6. 优先选择靠前的图片
                                score += (len(image_files) - image_files.index(img_path)) * 0.5
                                
                                print(f'[DEBUG] Image {img_path} score: {score:.1f}')
                                
                                if score > best_score:
                                    best_score = score
                                    best_photo = img_path
                                    
                            except Exception as e:
                                print(f'[DEBUG] Error analyzing image {img_path}: {e}')
                                continue
                        
                        # 如果找到最佳照片
                        if best_photo:
                            print(f'[DEBUG] Best photo selected: {best_photo} (score: {best_score:.1f})')
                            
                            # 确保photos目录存在
                            photos_dir = os.path.join('static', 'photos')
                            os.makedirs(photos_dir, exist_ok=True)
                            
                            # 生成唯一文件名
                            ext = os.path.splitext(best_photo)[1]
                            unique_filename = f"avatar_{uuid.uuid4().hex}{ext}"
                            save_path = os.path.join(photos_dir, unique_filename)
                            
                            # 提取并保存图片
                            with zf.open(best_photo) as f_in:
                                with open(save_path, 'wb') as f_out:
                                    f_out.write(f_in.read())
                            
                            # 生成网页可访问的路径
                            avatar_path = f"/static/photos/{unique_filename}"
                            print('[DEBUG] Photo extracted to:', avatar_path)
                        else:
                            # 尝试使用最后一张图片（可能是个人照片放在文档末尾）
                            if image_files:
                                last_image = image_files[-1]
                                print('[DEBUG] No best photo found, trying last image:', last_image)
                                
                                # 确保photos目录存在
                                photos_dir = os.path.join('static', 'photos')
                                os.makedirs(photos_dir, exist_ok=True)
                                
                                # 生成唯一文件名
                                ext = os.path.splitext(last_image)[1]
                                unique_filename = f"avatar_{uuid.uuid4().hex}{ext}"
                                save_path = os.path.join(photos_dir, unique_filename)
                                
                                # 提取并保存图片
                                with zf.open(last_image) as f_in:
                                    with open(save_path, 'wb') as f_out:
                                        f_out.write(f_in.read())
                                
                                # 生成网页可访问的路径
                                avatar_path = f"/static/photos/{unique_filename}"
                                print('[DEBUG] Photo extracted to:', avatar_path)
            except Exception as e:
                print('[DEBUG] Photo extraction failed:', str(e))
                import traceback
                traceback.print_exc()
            
            print('[DEBUG] DOCX parsing completed successfully')
            print('[DEBUG] Total extracted text length:', len(text))
            if len(text) > 0:
                print('[DEBUG] First 200 chars of text:', repr(text[:200]))
            else:
                print('[WARNING] No text extracted from DOCX file')
        except ImportError as e:
            print('[ERROR] docx module import failed:', str(e))
            print('[WARNING] DOCX parsing disabled. Please install python-docx module.')
            return None
        except Exception as e:
            print('[ERROR] DOCX parsing failed:', str(e))
            import traceback
            traceback.print_exc()
            return None
    else:
        print('[ERROR] Unsupported file type:', file_path)
        return None

    print('[DEBUG] Extracted text length:', len(text))
    if len(text) > 0:
        print('[DEBUG] First 500 chars of text:', text[:500])
    else:
        print('[WARNING] No text extracted from file')

    # 字段提取（支持多种简历格式）
    resume_data = {
        'name': '', 'job': '', 'intro': '', 'phone': '', 'email': '',
        'education': '', 'experience': '', 'skills': '', 'certificates': ''
    }

    # 提取姓名 - 支持更多格式
    name_patterns = [
        r'姓名[:：]\s*([^\s\n]{2,10})',
        r'Name[:：]\s*([^\s\n]{2,30})',
        r'^\s*([^\s\n]{2,10})\s*(?:男|女|\d+岁|\()',  # 开头可能的名字
        r'个人资料.*?姓名[:：]\s*([^\s\n]{2,10})',
        r'\n\s*([^\s\n]{2,10})\s*(?:先生|女士)',  # 名字+先生/女士
        r'求职意向.*?\n\s*([^\s\n]{2,10})',  # 求职意向附近的名字
        r'联系方式.*?\n\s*([^\s\n]{2,10})',  # 联系方式附近的名字
        r'简历.*?\n\s*([^\s\n]{2,10})',  # 简历标题附近的名字
        r'\n\s*([^\s\n]{2,10})\s*(?:\|\|/|\\)',  # 名字+分隔符
    ]
    for pattern in name_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            resume_data['name'] = match.group(1).strip()
            print('[DEBUG] Found name:', resume_data['name'])
            break

    # 提取求职意向
    job_patterns = [
        r'求职意向[:：]\s*(.+?)(?:\n|$)',
        r'应聘岗位[:：]\s*(.+?)(?:\n|$)',
        r'Position[:：]\s*(.+?)(?:\n|$)',
        r'意向岗位[:：]\s*(.+?)(?:\n|$)',
        r'求职目标[:：]\s*(.+?)(?:\n|$)',
        r'期望职位[:：]\s*(.+?)(?:\n|$)',
        r'目标职位[:：]\s*(.+?)(?:\n|$)',
        r'应聘职位[:：]\s*(.+?)(?:\n|$)',
        r'求职意向.*?\n\s*(.+?)(?:\n|$)',  # 更宽松的匹配
    ]
    for pattern in job_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            resume_data['job'] = match.group(1).strip()
            print('[DEBUG] Found job:', resume_data['job'])
            break

    # 提取联系电话
    phone_patterns = [
        r'电话[:：]\s*(\d{11}|\d{3,4}-?\d{7,8})',
        r'手机[:：]\s*(\d{11})',
        r'Phone[:：]\s*(\d{11}|\d{3,4}-?\d{7,8})',
        r'联系方式[:：].*?(\d{11})',
        r'[^\d](1[3-9]\d{9})[^\d]',  # 直接匹配11位手机号
        r'\b(1[3-9]\d{9})\b',  # 更宽松的手机号匹配
        r'Tel[:：]\s*(\d{11}|\d{3,4}-?\d{7,8})',
    ]
    for pattern in phone_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            resume_data['phone'] = match.group(1).strip()
            print('[DEBUG] Found phone:', resume_data['phone'])
            break

    # 提取邮箱 - 优先提取学校邮箱
    # 清理文本中的多余空格，特别是邮箱地址中的空格
    cleaned_text = text
    # 处理邮箱格式中的空格问题
    cleaned_text = re.sub(r'([a-zA-Z0-9._%+-])\s+([a-zA-Z0-9._%+-]@)', r'\1\2', cleaned_text)
    cleaned_text = re.sub(r'(@[a-zA-Z0-9.-]+)\s+([a-zA-Z])', r'\1\2', cleaned_text)
    cleaned_text = re.sub(r'(\.[a-zA-Z]{2,3})\s+(\.[a-zA-Z]{2,3})', r'\1\2', cleaned_text)
    
    print('[DEBUG] Cleaned text for email extraction (first 200 chars):', repr(cleaned_text[:200]))
    
    # 学校邮箱模式：匹配 .edu.cn、.edu、.ac.cn 等教育机构邮箱
    edu_email_patterns = [
        r'[\w\.-]+@(?:[\w-]+\.)?edu(?:\.[\w-]+)+',  # .edu.cn, .edu等
        r'[\w\.-]+@(?:[\w-]+\.)?ac(?:\.[\w-]+)+',   # .ac.cn等学术机构
        r'[\w\.-]+@[\w-]*\.(?:edu|ac)\.[a-z]{2,3}'   # 其他教育域名
    ]

    # 先尝试提取学校邮箱
    for pattern in edu_email_patterns:
        edu_email_match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if edu_email_match:
            resume_data['email'] = edu_email_match.group(0)
            print('[DEBUG] Found education email:', resume_data['email'])
            break

    # 如果没有找到学校邮箱，再提取普通邮箱
    if not resume_data['email']:
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', cleaned_text)
        if email_match:
            resume_data['email'] = email_match.group(0)
            print('[DEBUG] Found regular email:', resume_data['email'])

    # 提取教育经历 - 支持更多格式
    edu_patterns = [
        r'(教育经历|学历|Education)[:：]?\s*(.*?)(?=工作经历|项目经历|实习经历|技能|自我评价|专业技能|证书|荣誉|$)',
        r'(教育背景|学历背景)[:：]?\s*(.*?)(?=工作经历|项目经历|技能|证书|荣誉|$)',
        r'(教育|Education)[:：]?\s*(.*?)(?=工作|项目|技能|证书|$)',  # 更宽松的匹配
        r'(学习经历|academic)[:：]?\s*(.*?)(?=工作|项目|技能|证书|$)',  # 学习经历
    ]
    for pattern in edu_patterns:
        edu_section = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if edu_section:
            edu_text = edu_section.group(2).strip()
            print('[DEBUG] Education section found, length:', len(edu_text))

            # 从教育经历中提取学校邮箱
            if not resume_data['email']:
                edu_emails = re.findall(r'[\w\.-]+@(?:[\w-]+\.)?edu(?:\.[\w-]+)+', edu_text, re.IGNORECASE)
                if edu_emails:
                    resume_data['email'] = edu_emails[0]
                    print('[DEBUG] Found education email in education section:', resume_data['email'])

            # 改进的教育经历提取：支持多种格式
            edu_items = []

            # 尝试匹配包含时间段的教育经历
            time_pattern = r'(\d{4}[\./]?\d{0,2}[\./]?\d{0,2}?\s*[-~至到—]\s*\d{4}[\./]?\d{0,2}[\./]?\d{0,2}?)'

            # 按行分割并处理
            lines = [line.strip() for line in edu_text.split('\n') if line.strip()]
            current_edu = []

            for line in lines:
                # 如果包含时间段，可能是新的教育经历开始
                if re.search(time_pattern, line):
                    if current_edu:
                        edu_items.append(' '.join(current_edu))
                    current_edu = [line]
                else:
                    # 检查是否是新的教育经历（包含学校名称）
                    if any(keyword in line for keyword in ['大学', '学院', '学校', 'University', 'College', 'School', 'Institute']):
                        if current_edu:
                            edu_items.append(' '.join(current_edu))
                        current_edu = [line]
                    else:
                        # 继续当前教育经历
                        if current_edu:
                            current_edu.append(line)

            # 添加最后一条
            if current_edu:
                edu_items.append(' '.join(current_edu))

            # 如果没有找到结构化内容，尝试另一种方式提取
            if not edu_items:
                # 尝试匹配包含学校名称的段落
                school_pattern = r'([^\n]{20,200}(?:大学|学院|学校|University|College|School|Institute)[^\n]*)'
                edu_items = re.findall(school_pattern, edu_text)

            # 如果仍然没有找到，使用整个教育文本
            if not edu_items:
                edu_items = [edu_text[:500]]

            if edu_items:
                # 格式化教育经历，确保包含关键信息
                formatted_edu = []
                for item in edu_items[:5]:  # 最多保留5条
                    # 提取时间段
                    time_match = re.search(time_pattern, item)
                    # 提取学校名称
                    school_match = re.search(r'(.*?(?:大学|学院|学校|University|College|School|Institute).*?)(?:专业|学位|$)', item, re.IGNORECASE)
                    # 提取专业
                    major_match = re.search(r'(?:专业|Major|Majoring)[:：]?\s*([^\n]{5,50})', item, re.IGNORECASE)
                    # 提取学位
                    degree_match = re.search(r'(?:学位|Degree)[:：]?\s*([^\n]{2,20})', item, re.IGNORECASE)

                    parts = []
                    if time_match:
                        parts.append(time_match.group(1))
                    if school_match:
                        parts.append(school_match.group(1).strip())
                    if major_match:
                        parts.append(major_match.group(1).strip())
                    if degree_match:
                        parts.append(degree_match.group(1).strip())

                    if parts:
                        formatted_edu.append(' | '.join(parts))
                    else:
                        formatted_edu.append(item[:200])

                resume_data['education'] = '\n\n'.join(formatted_edu)
                print('[DEBUG] Found education items:', len(formatted_edu))
            else:
                resume_data['education'] = edu_text[:800]
            break

    # 提取工作/项目经历
    exp_patterns = [
        r'(工作经历|项目经历|实习经历|Experience)[:：]?\s*(.*?)(?=教育经历|学历|技能|自我评价|专业技能|证书|荣誉|$)',
        r'(工作背景|项目背景|实践经历)[:：]?\s*(.*?)(?=教育经历|技能|证书|荣誉|$)',
        r'(工作|项目|实习|Experience)[:：]?\s*(.*?)(?=教育|学历|技能|证书|$)',  # 更宽松的匹配
        r'(职业经历|work|project)[:：]?\s*(.*?)(?=教育|学历|技能|证书|$)',  # 职业经历
    ]
    for pattern in exp_patterns:
        exp_section = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if exp_section:
            exp_text = exp_section.group(2).strip()
            print('[DEBUG] Experience section found, length:', len(exp_text))
            
            # 尝试提取时间段+公司+职位格式
            exp_items = []
            
            # 按行分割并处理
            lines = [line.strip() for line in exp_text.split('\n') if line.strip()]
            current_exp = []
            
            # 时间段模式
            time_pattern = r'(\d{4}[\./]?\d{0,2}[\./]?\d{0,2}?\s*[-~至到—]\s*(?:\d{4}[\./]?\d{0,2}[\./]?\d{0,2}?|至今|现在))'
            
            for line in lines:
                # 如果包含时间段，可能是新的工作经历开始
                if re.search(time_pattern, line):
                    if current_exp:
                        exp_items.append(' '.join(current_exp))
                    current_exp = [line]
                else:
                    # 检查是否是新的工作经历（包含公司名称）
                    if any(keyword in line for keyword in ['公司', '企业', '集团', '公司', '有限公司', 'Corporation', 'Company', 'Ltd', 'Inc']):
                        if current_exp:
                            exp_items.append(' '.join(current_exp))
                        current_exp = [line]
                    else:
                        # 继续当前工作经历
                        if current_exp:
                            current_exp.append(line)
            
            # 添加最后一条
            if current_exp:
                exp_items.append(' '.join(current_exp))
            
            # 如果没有找到结构化内容，尝试另一种方式提取
            if not exp_items:
                # 尝试匹配包含公司名称的段落
                company_pattern = r'([^\n]{30,300}(?:公司|企业|集团|有限公司|Corporation|Company|Ltd|Inc)[^\n]*)'
                exp_items = re.findall(company_pattern, exp_text)
            
            # 如果仍然没有找到，使用整个工作经历文本
            if not exp_items:
                exp_items = [exp_text[:500]]
            
            if exp_items:
                resume_data['experience'] = '\n\n'.join([item[:300].strip() for item in exp_items[:3]])
                print('[DEBUG] Found experience items:', len(exp_items))
            else:
                resume_data['experience'] = exp_text[:800]
            break

    # 提取技能
    skill_patterns = [
        r'(技能|专业技能|Skills)[:：]?\s*(.*?)(?=教育经历|工作经历|自我评价|证书|荣誉|$)',
        r'(技术栈|专长|技术能力)[:：]?\s*(.*?)(?=教育经历|工作经历|证书|荣誉|$)',
        r'(技能|Skills)[:：]?\s*(.*?)(?=教育|工作|证书|$)',  # 更宽松的匹配
        r'(技术|技术能力|tech|skills)[:：]?\s*(.*?)(?=教育|工作|证书|$)',  # 技术能力
    ]
    for pattern in skill_patterns:
        skill_section = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if skill_section:
            skills_text = skill_section.group(2).strip()[:500]
            # 清理技能文本，去除多余空白和特殊字符
            skills_text = re.sub(r'\s+', ' ', skills_text)
            # 将技能文本转换为逗号分隔的格式
            # 尝试匹配常见的技能分隔符
            skills = re.split(r'[，,；;\s]+', skills_text)
            # 过滤空技能和太短的技能
            skills = [skill.strip() for skill in skills if skill.strip() and len(skill.strip()) > 1]
            # 去重
            skills = list(set(skills))
            # 限制技能数量
            skills = skills[:20]
            resume_data['skills'] = ','.join(skills)
            print('[DEBUG] Found skills:', resume_data['skills'])
            break

    # 提取证书/荣誉
    cert_patterns = [
        r'(证书|荣誉|奖项|Certificates|Awards)[:：]?\s*(.*?)(?=教育经历|工作经历|技能|$)',
        r'(证书荣誉|获奖情况)[:：]?\s*(.*?)(?=教育经历|工作经历|技能|$)',
        r'(证书|荣誉|奖项)[:：]?\s*(.*?)(?=教育|工作|技能|$)',  # 更宽松的匹配
        r'(资格证书|证书奖项|certification|honor)[:：]?\s*(.*?)(?=教育|工作|技能|$)',  # 资格证书
    ]
    for pattern in cert_patterns:
        cert_section = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if cert_section:
            resume_data['certificates'] = cert_section.group(2).strip()[:500]
            print('[DEBUG] Found certificates')
            break

    # 提取个人简介/自我评价
    intro_patterns = [
        r'(个人简介|自我评价|自我描述|Self Introduction)[:：]?\s*(.*?)(?=教育经历|工作经历|技能|证书|$)',
        r'(个人介绍|自我介绍)[:：]?\s*(.*?)(?=教育经历|工作经历|技能|证书|$)',
        r'(个人简介|自我评价)[:：]?\s*(.*?)(?=教育|工作|技能|证书|$)',  # 更宽松的匹配
        r'(简介|自我介绍|personal|summary)[:：]?\s*(.*?)(?=教育|工作|技能|证书|$)',  # 简介
    ]
    for pattern in intro_patterns:
        intro_section = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if intro_section:
            resume_data['intro'] = intro_section.group(2).strip()[:500]
            print('[DEBUG] Found intro')
            break

    # 添加提取的头像路径
    if avatar_path:
        resume_data['avatar'] = avatar_path
        print('[DEBUG] Added avatar path to result')

    # 过滤空值
    result = {k: v for k, v in resume_data.items() if v}
    print('[DEBUG] Final parsed data:', result)
    
    # 如果没有提取到任何信息，返回None
    if not result:
        print('[DEBUG] No information extracted')
        return None
    
    return result