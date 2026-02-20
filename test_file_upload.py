#!/usr/bin/env python3
# 测试文件上传功能的调试脚本

import os
import sys
import requests
import uuid

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_file_upload():
    """测试文件上传功能"""
    print("=== 测试文件上传功能 ===")
    
    # 测试服务器地址
    server_url = "http://localhost:5000"
    
    # 首先登录
    print("\n1. 测试登录...")
    login_data = {
        "username": "admin",
        "password": "your_new_password"
    }
    
    session = requests.Session()
    login_response = session.post(f"{server_url}/login", data=login_data)
    print(f"登录状态码: {login_response.status_code}")
    print(f"登录结果: {'成功' if login_response.status_code == 200 else '失败'}")
    
    # 检查是否有测试用的Word文档
    test_docx_path = None
    for root, dirs, files in os.walk('.'):
        for file in files:
            if file.endswith('.docx') and 'test' in file.lower():
                test_docx_path = os.path.join(root, file)
                break
        if test_docx_path:
            break
    
    if not test_docx_path:
        # 创建一个简单的测试Word文档
        print("\n2. 创建测试Word文档...")
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('测试简历', 0)
            doc.add_paragraph('姓名: 测试用户')
            doc.add_paragraph('电话: 13800138000')
            doc.add_paragraph('邮箱: test@example.com')
            doc.add_heading('教育经历', level=1)
            doc.add_paragraph('2020-2024 | 北京大学 | 计算机科学与技术')
            doc.add_heading('工作经历', level=1)
            doc.add_paragraph('2024-至今 | 腾讯科技 | 软件工程师')
            doc.add_heading('技能', level=1)
            doc.add_paragraph('Python, Java, JavaScript')
            
            test_docx_path = f"test_resume_{uuid.uuid4().hex}.docx"
            doc.save(test_docx_path)
            print(f"创建测试文档成功: {test_docx_path}")
            print(f"文件大小: {os.path.getsize(test_docx_path)} bytes")
        except Exception as e:
            print(f"创建测试文档失败: {e}")
            return
    
    print(f"\n3. 测试文件: {test_docx_path}")
    print(f"文件大小: {os.path.getsize(test_docx_path)} bytes")
    
    # 测试文件上传
    print("\n4. 测试文件上传...")
    try:
        with open(test_docx_path, 'rb') as f:
            files = {
                'resume_file': (os.path.basename(test_docx_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            data = {
                'action': 'upload'
            }
            
            print("发送上传请求...")
            upload_response = session.post(f"{server_url}/edit", files=files, data=data)
            
            print(f"上传状态码: {upload_response.status_code}")
            print(f"上传结果: {'成功' if upload_response.status_code == 200 else '失败'}")
            
            # 检查响应内容
            if upload_response.status_code == 200:
                print("\n5. 检查上传的文件...")
                # 列出服务器上的documents文件夹内容
                documents_dir = os.path.join('static', 'documents')
                if os.path.exists(documents_dir):
                    docx_files = [f for f in os.listdir(documents_dir) if f.endswith('.docx')]
                    print(f"服务器上的docx文件数量: {len(docx_files)}")
                    
                    # 检查最新上传的文件
                    if docx_files:
                        # 按修改时间排序
                        docx_files.sort(key=lambda x: os.path.getmtime(os.path.join(documents_dir, x)), reverse=True)
                        latest_file = docx_files[0]
                        latest_file_path = os.path.join(documents_dir, latest_file)
                        
                        print(f"最新上传的文件: {latest_file}")
                        print(f"文件大小: {os.path.getsize(latest_file_path)} bytes")
                        
                        # 尝试读取文件内容
                        try:
                            from docx import Document
                            doc = Document(latest_file_path)
                            print(f"段落数量: {len(doc.paragraphs)}")
                            print("\n文件内容:")
                            for i, paragraph in enumerate(doc.paragraphs):
                                text = paragraph.text.strip()
                                if text:
                                    print(f"段落 {i+1}: {text}")
                        except Exception as e:
                            print(f"读取文件内容失败: {e}")
                else:
                    print(f"服务器documents目录不存在: {documents_dir}")
                    
    except Exception as e:
        print(f"上传测试失败: {e}")
        import traceback
        traceback.print_exc()
    
    # 清理测试文件
    if os.path.exists(test_docx_path) and 'test_resume_' in test_docx_path:
        try:
            os.remove(test_docx_path)
            print(f"\n清理测试文件: {test_docx_path}")
        except:
            pass
    
    print("\n=== 测试完成 ===")

if __name__ == "__main__":
    test_file_upload()