#!/usr/bin/env python3
# 只测试解析逻辑的简单脚本

import os
import sys

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_parse_only():
    """只测试解析逻辑"""
    print("=== 只测试解析逻辑 ===")
    
    # 测试文件路径
    test_files = []
    documents_dir = os.path.join('static', 'documents')
    
    if os.path.exists(documents_dir):
        # 获取最近的几个docx文件
        docx_files = [f for f in os.listdir(documents_dir) if f.endswith('.docx')]
        if docx_files:
            # 按修改时间排序
            docx_files.sort(key=lambda x: os.path.getmtime(os.path.join(documents_dir, x)), reverse=True)
            # 测试前3个文件
            for i, docx_file in enumerate(docx_files[:3]):
                file_path = os.path.join(documents_dir, docx_file)
                test_files.append((f"文件{i+1}: {docx_file}", file_path))
    
    # 如果没有docx文件，创建一个测试文件
    if not test_files:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('测试简历', 0)
            doc.add_paragraph('姓名: 测试用户')
            doc.add_paragraph('电话: 13800138000')
            doc.add_paragraph('邮箱: test@example.com')
            doc.add_heading('教育经历', level=1)
            doc.add_paragraph('2020-2024 | 北京大学 | 计算机科学与技术')
            doc.add_heading('工作经历', level=1)
            doc.add_paragraph('2024-至今 | 腾讯科技 | 软件工程师')
            doc.add_heading('技能', level=1)
            doc.add_paragraph('Python, Java, JavaScript')
            
            test_file = 'test_parse_resume.docx'
            doc.save(test_file)
            test_files.append(("测试文件", test_file))
            print(f"创建测试文件: {test_file}")
        except Exception as e:
            print(f"创建测试文件失败: {e}")
            return
    
    # 导入解析函数
    try:
        from parse_resume_optimized import parse_resume_optimized as parse_resume
        print("\n解析函数导入成功")
    except Exception as e:
        print(f"解析函数导入失败: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # 测试每个文件
    for file_desc, file_path in test_files:
        print(f"\n=== 测试 {file_desc} ===")
        print(f"文件路径: {file_path}")
        print(f"文件存在: {os.path.exists(file_path)}")
        
        if os.path.exists(file_path):
            print(f"文件大小: {os.path.getsize(file_path)} bytes")
            
            # 先检查文件是否为有效的Word文档
            try:
                from docx import Document
                doc = Document(file_path)
                print(f"Word文档加载成功")
                print(f"段落数量: {len(doc.paragraphs)}")
                
                # 打印前几个段落内容
                print("前5个段落内容:")
                for i, paragraph in enumerate(doc.paragraphs[:5]):
                    text = paragraph.text.strip()
                    if text:
                        print(f"  段落 {i+1}: {text}")
            except Exception as e:
                print(f"Word文档加载失败: {e}")
                continue
            
            # 测试解析
            print("\n测试解析...")
            try:
                # 重定向标准输出到文件，避免截断
                import io
                old_stdout = sys.stdout
                sys.stdout = io.StringIO()
                
                result = parse_resume(file_path)
                
                # 恢复标准输出并获取输出内容
                output = sys.stdout.getvalue()
                sys.stdout = old_stdout
                
                print("解析输出:")
                print(output)
                print(f"解析结果: {result}")
                
                if result:
                    print("解析成功！提取的字段:")
                    for key, value in result.items():
                        print(f"  {key}: {value}")
                else:
                    print("解析失败，未提取到有效信息")
                    
            except Exception as e:
                print(f"解析失败: {e}")
                import traceback
                traceback.print_exc()
        else:
            print("文件不存在")
    
    # 清理测试文件
    if os.path.exists('test_parse_resume.docx'):
        try:
            os.remove('test_parse_resume.docx')
            print(f"\n清理测试文件")
        except:
            pass
    
    print("\n=== 测试完成 ===")

if __name__ == "__main__":
    test_parse_only()