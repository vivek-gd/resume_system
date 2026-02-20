# 测试docx模块的导入和使用
print('开始测试docx模块的导入和使用...')

try:
    print('尝试导入docx模块...')
    from docx import Document
    print('✓ docx模块导入成功')
    
    # 测试创建Document对象
    print('尝试创建Document对象...')
    doc = Document()
    print('✓ Document对象创建成功')
    
    # 测试添加段落
    print('尝试添加段落...')
    doc.add_paragraph('测试段落')
    print('✓ 段落添加成功')
    
    # 测试保存文档
    print('尝试保存文档...')
    doc.save('test_docx_import.docx')
    print('✓ 文档保存成功')
    
    # 测试打开现有文档
    print('尝试打开现有文档...')
    doc = Document('test_resume_valid.docx')
    print('✓ 现有文档打开成功')
    print('文档包含', len(doc.paragraphs), '个段落')
    
    # 测试读取段落内容
    print('尝试读取段落内容...')
    for i, para in enumerate(doc.paragraphs):
        print(f'段落 {i+1}: {para.text}')
    print('✓ 段落内容读取成功')
    
    print('\n测试完成，docx模块工作正常！')
except Exception as e:
    print('✗ 错误:', str(e))
    import traceback
    traceback.print_exc()