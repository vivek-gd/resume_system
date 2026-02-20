#!/usr/bin/env python3
# 直接测试文件处理逻辑的调试脚本

import os
import sys
import uuid

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_file_handling():
    """直接测试文件处理逻辑"""
    print("=== 测试文件处理逻辑 ===")
    
    # 创建测试Word文档
    print("\n1. 创建测试Word文档...")
    test_docx_path = None
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('测试简历', 0)
        doc.add_paragraph('姓名: 测试用户')
        doc.add_paragraph('电话: 13800138000')
        doc.add_paragraph('邮箱: test@example.com')
        doc.add_heading('教育经历', level=1)
        doc.add_paragraph('2020-2024 | 北京大学 | 计算机科学与技术')
        doc.add_heading('工作经历', level=1)
        doc.add_paragraph('2024-至今 | 腾讯科技 | 软件工程师')
        doc.add_heading('技能', level=1)
        doc.add_paragraph('Python, Java, JavaScript')
        
        test_docx_path = f"test_resume_{uuid.uuid4().hex}.docx"
        doc.save(test_docx_path)
        print(f"创建测试文档成功: {test_docx_path}")
        print(f"文件大小: {os.path.getsize(test_docx_path)} bytes")
        
        # 验证文档内容
        doc_check = Document(test_docx_path)
        print(f"文档段落数量: {len(doc_check.paragraphs)}")
        print("文档内容:")
        for i, paragraph in enumerate(doc_check.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"  段落 {i+1}: {text}")
                
    except Exception as e:
        print(f"创建测试文档失败: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # 模拟文件保存逻辑
    print("\n2. 测试文件保存逻辑...")
    try:
        # 模拟Flask的文件保存逻辑
        upload_folder = 'static'
        documents_dir = os.path.join(upload_folder, 'documents')
        os.makedirs(documents_dir, exist_ok=True)
        
        # 生成唯一文件名
        ext = os.path.splitext(test_docx_path)[1]
        unique_filename = f"{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(documents_dir, unique_filename)
        
        print(f"目标保存路径: {file_path}")
        
        # 复制文件
        import shutil
        shutil.copy2(test_docx_path, file_path)
        
        print(f"文件复制成功")
        print(f"保存后文件大小: {os.path.getsize(file_path)} bytes")
        
        # 验证保存的文件
        doc_saved = Document(file_path)
        print(f"保存后文档段落数量: {len(doc_saved.paragraphs)}")
        print("保存后文档内容:")
        for i, paragraph in enumerate(doc_saved.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"  段落 {i+1}: {text}")
                
    except Exception as e:
        print(f"文件保存测试失败: {e}")
        import traceback
        traceback.print_exc()
    
    # 测试解析逻辑
    print("\n3. 测试简历解析逻辑...")
    try:
        from parse_resume_optimized import parse_resume_optimized as parse_resume
        
        # 解析测试文件
        result = parse_resume(test_docx_path)
        print(f"解析结果: {result}")
        
        if result:
            print("解析成功！提取的字段:")
            for key, value in result.items():
                print(f"  {key}: {value}")
        else:
            print("解析失败，未提取到有效信息")
            
    except Exception as e:
        print(f"解析测试失败: {e}")
        import traceback
        traceback.print_exc()
    
    # 清理测试文件
    if test_docx_path and os.path.exists(test_docx_path):
        try:
            os.remove(test_docx_path)
            print(f"\n清理测试文件: {test_docx_path}")
        except:
            pass
    
    print("\n=== 测试完成 ===")

if __name__ == "__main__":
    test_file_handling()