#!/usr/bin/env python3
# 测试documents文件夹中的实际docx文件

import os
import sys

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_actual_docx():
    """测试documents文件夹中的实际docx文件"""
    print("=== 测试documents文件夹中的实际docx文件 ===")
    
    # 检查documents目录
    documents_dir = os.path.join('static', 'documents')
    print(f"检查目录: {documents_dir}")
    print(f"目录存在: {os.path.exists(documents_dir)}")
    
    if not os.path.exists(documents_dir):
        print("目录不存在")
        return
    
    # 获取所有docx文件
    docx_files = [f for f in os.listdir(documents_dir) if f.endswith('.docx')]
    print(f"找到 {len(docx_files)} 个docx文件")
    
    if not docx_files:
        print("没有找到docx文件")
        return
    
    # 测试每个docx文件
    for i, docx_file in enumerate(docx_files[:5]):  # 只测试前5个
        file_path = os.path.join(documents_dir, docx_file)
        print(f"\n=== 测试文件 {i+1}: {docx_file} ===")
        print(f"文件路径: {file_path}")
        print(f"文件大小: {os.path.getsize(file_path)} bytes")
        
        # 尝试使用python-docx打开
        try:
            from docx import Document
            doc = Document(file_path)
            print(f"✓ python-docx打开成功")
            print(f"段落数量: {len(doc.paragraphs)}")
            
            # 检查段落内容
            print("段落内容:")
            for j, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    print(f"  段落 {j+1}: '{text}'")
                else:
                    print(f"  段落 {j+1}: [空]")
            
            # 检查是否只有一个段落且内容为"0"
            if len(doc.paragraphs) == 1:
                first_paragraph = doc.paragraphs[0].text.strip()
                if first_paragraph == "0":
                    print("⚠️  发现问题：文件只包含一个段落，内容为'0'")
                    
                    # 尝试使用其他方法读取文件
                    print("\n尝试使用zipfile检查文件结构...")
                    try:
                        import zipfile
                        with zipfile.ZipFile(file_path, 'r') as zf:
                            print(f"✓ 文件是有效的ZIP文件")
                            print("文件内容:")
                            for name in zf.namelist():
                                print(f"  - {name}")
                            
                            # 尝试读取document.xml
                            if 'word/document.xml' in zf.namelist():
                                print("\n读取document.xml...")
                                with zf.open('word/document.xml') as f:
                                    xml_content = f.read().decode('utf-8')
                                    print(f"XML长度: {len(xml_content)} 字符")
                                    print(f"前500字符: {xml_content[:500]}...")
                    except Exception as e:
                        print(f"✗ zipfile检查失败: {e}")
                        
        except Exception as e:
            print(f"✗ python-docx打开失败: {e}")
            import traceback
            traceback.print_exc()
        
        # 测试解析函数
        print("\n测试解析函数...")
        try:
            from parse_resume_optimized import parse_resume_optimized as parse_resume
            result = parse_resume(file_path)
            print(f"解析结果: {result}")
        except Exception as e:
            print(f"解析失败: {e}")
            import traceback
            traceback.print_exc()
    
    print("\n=== 测试完成 ===")

if __name__ == "__main__":
    test_actual_docx()