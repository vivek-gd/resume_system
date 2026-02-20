#!/usr/bin/env python3
# 测试特定的docx文件

import os
import sys

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_specific_file():
    """测试特定的docx文件"""
    print("=== 测试特定的docx文件 ===")
    
    # 测试用户提到的文件
    test_file = "5d44f5537ce7466d8dfee1a15f42c432.docx"
    documents_dir = os.path.join('static', 'documents')
    file_path = os.path.join(documents_dir, test_file)
    
    print(f"测试文件: {test_file}")
    print(f"文件路径: {file_path}")
    print(f"文件存在: {os.path.exists(file_path)}")
    
    if not os.path.exists(file_path):
        print("文件不存在，测试其他docx文件...")
        # 查找其他docx文件
        if os.path.exists(documents_dir):
            docx_files = [f for f in os.listdir(documents_dir) if f.endswith('.docx')]
            if docx_files:
                test_file = docx_files[0]
                file_path = os.path.join(documents_dir, test_file)
                print(f"使用文件: {test_file}")
                print(f"文件路径: {file_path}")
            else:
                print("没有找到docx文件")
                return
        else:
            print("documents目录不存在")
            return
    
    print(f"文件大小: {os.path.getsize(file_path)} bytes")
    
    # 1. 尝试使用python-docx打开
    print("\n1. 使用python-docx打开文件...")
    try:
        from docx import Document
        doc = Document(file_path)
        print(f"✓ 打开成功")
        print(f"段落数量: {len(doc.paragraphs)}")
        
        print("段落内容:")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            print(f"  段落 {i+1}: '{text}'")
            
        # 2. 尝试使用zipfile检查文件结构
        print("\n2. 使用zipfile检查文件结构...")
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zf:
                print(f"✓ 是有效的ZIP文件")
                print("文件内容:")
                for name in zf.namelist():
                    print(f"  - {name}")
                
                # 3. 尝试读取document.xml
                print("\n3. 读取document.xml...")
                if 'word/document.xml' in zf.namelist():
                    with zf.open('word/document.xml') as f:
                        xml_content = f.read().decode('utf-8', errors='ignore')
                        print(f"XML长度: {len(xml_content)} 字符")
                        print(f"前1000字符: {xml_content[:1000]}...")
                else:
                    print("✗ 没有找到word/document.xml")
                    
        except Exception as e:
            print(f"✗ zipfile检查失败: {e}")
            import traceback
            traceback.print_exc()
            
        # 4. 测试解析函数
        print("\n4. 测试解析函数...")
        try:
            from parse_resume_optimized import parse_resume_optimized as parse_resume
            # 重定向输出到文件
            import io
            old_stdout = sys.stdout
            sys.stdout = io.StringIO()
            
            result = parse_resume(file_path)
            
            # 恢复输出
            output = sys.stdout.getvalue()
            sys.stdout = old_stdout
            
            print("解析输出:")
            print(output)
            print(f"解析结果: {result}")
            
        except Exception as e:
            print(f"✗ 解析失败: {e}")
            import traceback
            traceback.print_exc()
            
    except Exception as e:
        print(f"✗ python-docx打开失败: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n=== 测试完成 ===")

if __name__ == "__main__":
    test_specific_file()