# 使用python-docx创建测试DOCX文件
from docx import Document

print('开始创建测试DOCX文件...')

# 创建一个新的DOCX文档
doc = Document()

# 添加标题
doc.add_heading('测试简历', level=1)

# 添加姓名和求职意向
doc.add_paragraph('姓名：张三')
doc.add_paragraph('求职意向：软件工程师')

# 添加联系方式
doc.add_heading('联系方式', level=2)
doc.add_paragraph('电话：13812345678')
doc.add_paragraph('邮箱：zhangsan@example.com')

# 添加教育经历
doc.add_heading('教育经历', level=2)
doc.add_paragraph('2018-2022 北京大学 计算机科学与技术 本科')

# 添加工作经历
doc.add_heading('工作经历', level=2)
doc.add_paragraph('2022-至今 腾讯科技有限公司 软件工程师')

# 添加技能
doc.add_heading('技能', level=2)
doc.add_paragraph('Python, Java, C++, JavaScript, HTML, CSS, React, Vue, MySQL, MongoDB')

# 添加证书
doc.add_heading('证书', level=2)
doc.add_paragraph('英语四级')
doc.add_paragraph('英语六级')

# 保存文档
doc.save('test_resume_valid.docx')
print('测试DOCX文件创建成功：test_resume_valid.docx')