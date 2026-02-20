#!/usr/bin/env python3
# 直接测试Word文档解析的调试脚本

import os
import sys

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_docx_parsing(file_path):
    """测试Word文档解析"""
    print(f"\n=== 测试Word文档解析: {file_path} ===")
    print(f"文件存在: {os.path.exists(file_path)}")
    if os.path.exists(file_path):
        print(f"文件大小: {os.path.getsize(file_path)} bytes")
    
    try:
        # 尝试导入docx模块
        from docx import Document
        print("docx模块导入成功")
        
        # 加载文档
        doc = Document(file_path)
        print("文档加载成功")
        print(f"段落数量: {len(doc.paragraphs)}")
        print(f"表格数量: {len(doc.tables)}")
        print(f"章节数量: {len(doc.sections)}")
        
        # 测试段落提取
        print("\n=== 段落内容 ===")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"段落 {i+1}: {text[:100]}..." if len(text) > 100 else f"段落 {i+1}: {text}")
            else:
                print(f"段落 {i+1}: [空]")
        
        # 测试表格提取
        print("\n=== 表格内容 ===")
        for i, table in enumerate(doc.tables):
            print(f"表格 {i+1}: {len(table.rows)}行 × {len(table.columns)}列")
            for j, row in enumerate(table.rows):
                row_text = []
                for k, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_text.append(cell_text if cell_text else "[空]")
                print(f"  行 {j+1}: {' | '.join(row_text)}")
        
        # 测试XML直接提取
        print("\n=== XML直接提取 ===")
        try:
            import xml.etree.ElementTree as ET
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            
            print("尝试遍历文档部分...")
            for part in doc.part.iter_parts():
                print(f"部分类型: {part.content_type}")
                if part.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml':
                    print("找到主文档XML")
                    try:
                        root = ET.fromstring(part.blob)
                        print("XML解析成功")
                        # 查找所有文本节点
                        text_nodes = root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        print(f"找到 {len(text_nodes)} 个文本节点")
                        for i, text_node in enumerate(text_nodes[:10]):  # 只显示前10个
                            if text_node.text:
                                print(f"  文本节点 {i+1}: {repr(text_node.text)}")
                    except Exception as e:
                        print(f"XML解析失败: {e}")
        except Exception as e:
            print(f"XML提取失败: {e}")
        
        # 测试另一种方法
        print("\n=== 测试python-docx的其他方法 ===")
        try:
            # 测试文档属性
            print(f"文档标题: {doc.core_properties.title}")
            print(f"文档作者: {doc.core_properties.author}")
            print(f"文档主题: {doc.core_properties.subject}")
        except Exception as e:
            print(f"属性读取失败: {e}")
            
    except ImportError as e:
        print(f"docx模块导入失败: {e}")
    except Exception as e:
        print(f"解析失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # 测试documents文件夹中的所有docx文件
    documents_dir = os.path.join(os.path.dirname(__file__), 'static', 'documents')
    print(f"检查目录: {documents_dir}")
    
    if os.path.exists(documents_dir):
        docx_files = [f for f in os.listdir(documents_dir) if f.endswith('.docx')]
        print(f"找到 {len(docx_files)} 个docx文件:")
        for docx_file in docx_files:
            print(f"  - {docx_file}")
        
        # 测试每个文件
        for docx_file in docx_files:
            file_path = os.path.join(documents_dir, docx_file)
            test_docx_parsing(file_path)
    else:
        print(f"目录不存在: {documents_dir}")
    
    print("\n=== 测试完成 ===")